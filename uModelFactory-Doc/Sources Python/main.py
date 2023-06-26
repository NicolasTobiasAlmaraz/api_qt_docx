from xmlreader import XmlReader
from writerodt import OdtWriter

from docx import Document
from docx.shared import Cm

readerXml = XmlReader("genOdt.xml")
path = readerXml.getOutputPath()
writer = OdtWriter(path)
while readerXml.hasItems():
    (item, typeItem) = readerXml.getItem()

    if typeItem == "pageSetUp":
        writer.setPageConfig(item)

    if typeItem == "text":
        writer.writeText(item)

    if typeItem == "image":
        writer.writeImage(item)

    if typeItem == "table":
        writer.writeTable(item)

    if typeItem == "list":
        writer.writeList(item)

    if typeItem == "header":
        writer.writeHeader(item)

    if typeItem == "footer":
        writer.writeFooter(item)

writer.closeFile()

###########################################################
###########################################################

# from docx import Document
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml
# from lxml import etree
#
# def escribir_encabezado(documento, lista_parrafos):
#     section = documento.sections[0]
#     header = section.header
#
#     for parrafo in lista_parrafos:
#         xml = parrafo._element.xml
#         tree = etree.fromstring(xml)
#         header._element.append(tree)
#
# def escribir_pie_de_pagina(documento, lista_parrafos):
#     section = documento.sections[0]
#     footer = section.footer
#
#     for parrafo in lista_parrafos:
#         xml = parrafo._element.xml
#         tree = etree.fromstring(xml)
#         footer._element.append(tree)
#
# # Ejemplo de uso
# documento = Document()
#
# parrafos_encabezado = [
#     Document().add_paragraph("Este es el encabezado"),
#     Document().add_paragraph("Línea 2 del encabezado"),
#     Document().add_paragraph("Línea 3 del encabezado")
# ]
#
# parrafos_pie_de_pagina = [
#     Document().add_paragraph("Este es el pie de página"),
#     Document().add_paragraph("Línea 2 del pie de página"),
#     Document().add_paragraph("Línea 3 del pie de página")
# ]
#
# escribir_encabezado(documento, parrafos_encabezado)
# escribir_pie_de_pagina(documento, parrafos_pie_de_pagina)
#
# documento.save("documento_con_encabezado_y_pie.docx")


###########################################################
###########################################################

#
# from docx import Document
# from docx.shared import Inches, Pt
# document = Document()
# header = document.sections[0].header
# htable=header.add_table(1, 2, Inches(6))
# htab_cells=htable.rows[0].cells
# ht0=htab_cells[0].add_paragraph()
# kh=ht0.add_run()
# kh.add_picture('D:/UTN/uModelFactory/uModelFactory-Documentacion/ProyectoCosechadoraAgro/doc_CosechadoraAgro/logoUTN.png', width=Inches(1))
# ht1=htab_cells[1].add_paragraph('Texto1')
# ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#
# ht1=htab_cells[1].add_paragraph('Texto2')
# ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#
# document.save('yourdoc.docx')
#
