# Xml items
from pagedocx import pageDocx
from tabledocx import tableDocx
from textdocx import textDocx
from imagedocx import imageDocx
from listdocx import listDocx

from docx import Document
from docx.shared import  RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt, Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def separarRgb(cadena):
    componentes = [int(cadena[i:i+2], 16) for i in range(0, len(cadena), 2)]
    return componentes



class DocxWriter:
    def __init__(self, path):
        self.listElementsHeader = []
        self.listElementsFooter = []
        self.path = path
        self.myTextDoc = Document()
        self.cantStyles = 0
        self.cantTables = 0
        self.aux = False

    # Genera un objeto texto listo para agregar al documento
    def generateText(self, text: textDocx, objDestino):
        paragraph = objDestino.add_paragraph()

        # Reajusto texto teniendo en cuenta los \t
        if text.mContent is not None:
            newText = ""
            cantSpaces = 0
            for c in text.mContent:
                if c == " ":
                    cantSpaces += 1
                    if cantSpaces == 4:
                        newText += "\t"
                        cantSpaces = 0
                else:
                    if cantSpaces != 4 and cantSpaces != 0:
                        newText += " "
                    cantSpaces = 0
                    newText += c
            text.mContent = newText

        # Establecer el contenido del párrafo
        content = text.mContent
        paragraph.add_run(content)

        # Aplicar formatos
        run = paragraph.runs[0]

        if text.mBold == "1":
            run.bold = True

        if text.mUnderline == "1":
            run.underline = True

        if text.mItalic == "1":
            run.italic = True

        if text.mStrikethrough == "1":
            run.strike = True

        # Establecer alineación
        alignment = text.mAlign.lower()
        paragraph.alignment = Document().styles['Normal'].paragraph_format.alignment
        if alignment == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'justify':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Establecer color de fuente
        rgb = separarRgb(text.mFillColor)
        run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

        # Establecer color de fondo
        #rgb = separarRgb(text.mBackgroundColor)
        #paragraph.runs[0].font.highlight_color = RGBColor(rgb[0], rgb[1], rgb[2])

        # Establecer fuente y tamaño
        run.font.name = text.mFont
        run.font.size = Pt(int(text.mLetterSize))

        return paragraph

    # Genera un objeto imagen listo para agregar al documento
    def generateImage(self, image: imageDocx, objDestino):

        paragraph = objDestino.add_paragraph()

        # Insertar la imagen con tamaño personalizado
        run = paragraph.add_run()
        picture = run.add_picture(image.mImgPath)
        picture.width = Pt(int(image.mWidth)*2/3)
        picture.height = Pt(int(image.mHeight)*2/3)

        # Establecer el anclaje de la imagen
        if image.mAnchor == 'paralel':
            picture.inline = False
        else:
            picture.inline = True

        # Establecer la alineación del párrafo
        if image.mAlign == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif image.mAlign == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif image.mAlign == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        return paragraph

    # Stea formato
    def setPageConfig(self, setUp: pageDocx):
        seccion = self.myTextDoc.sections[0]

        # Establecer el tamaño de página
        seccion.page_width = Mm(int(setUp.mSizeWidth))
        seccion.page_height = Mm(int(setUp.mSizeHeight))

        # Establecer los márgenes de página
        seccion.left_margin = Mm(int(setUp.mMarginLeft))
        seccion.right_margin = Mm(int(setUp.mMarginRight))
        seccion.top_margin = Mm(int(setUp.mMarginTop))
        seccion.bottom_margin = Mm(int(setUp.mMarginBottom))

    # Escribe Footer
    def writeFooter(self, footer):
        self.listElementsFooter = footer
        fText = False
        fImg = False
        col = 1
        colImg = 0
        colTxt = 0
        alignTxt = WD_PARAGRAPH_ALIGNMENT.CENTER
        alignImg = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Me fijo que tipo de elementos tengo
        for elem in footer:
            if elem["type"] == "text":
                fText = True
            if elem["type"] == "image":
                fImg = True

        # Si tengo texto e imagenes hago dos columnas
        if fText and fImg:
            col = 2
            colImg = 0
            colTxt = 1
            alignTxt = WD_PARAGRAPH_ALIGNMENT.RIGHT
            alignImg = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Si no tengo nada slgo
        if fText == False and fImg == False:
            return

        # Divido en las columnas que corresponda
        footer = self.myTextDoc.sections[0].footer
        htable = footer.add_table(1, col, Inches(6))
        htab_cells = htable.rows[0].cells

        # Recorro hasta que no haya mas elementos
        while len(self.listElementsFooter) != 0:
            item = self.listElementsFooter.pop(0)
            if item["type"] == "text":
                text = item["item"].mContent
                htTxt = htab_cells[colTxt].add_paragraph(text)
                htTxt.alignment = alignTxt

            if item["type"] == "image":
                htImg = htab_cells[colImg].add_paragraph()
                kh = htImg.add_run()
                path = item["item"].mImgPath
                width = item["item"].mWidth
                kh.add_picture(path, width=Inches(2))

    # Escribe Header
    def writeHeader(self, header):
        self.listElementsHeader = header
        fText = False
        fImg = False
        col = 1
        colImg = 0
        colTxt = 0
        alignTxt = WD_PARAGRAPH_ALIGNMENT.CENTER
        alignImg = WD_PARAGRAPH_ALIGNMENT.CENTER
        for elem in header:
            if elem["type"] == "text":
                fText = True
            if elem["type"] == "image":
                fImg = True

        if fText and fImg:
            col = 2
            colTxt = 1
            colImg = 0
            alignTxt = WD_PARAGRAPH_ALIGNMENT.RIGHT
            alignImg = WD_PARAGRAPH_ALIGNMENT.LEFT

        if fText==False and fImg==False:
            return

        header = self.myTextDoc.sections[0].header
        htable = header.add_table(1, col, Inches(6))
        htab_cells = htable.rows[0].cells

        while len(self.listElementsHeader) != 0:
            item = self.listElementsHeader.pop(0)
            if item["type"] == "text":
                text = item["item"].mContent
                htTxt = htab_cells[colTxt].add_paragraph(text)
                htTxt.alignment = alignTxt

            if item["type"] == "image":
                ht0 = htab_cells[colImg].add_paragraph()
                kh = ht0.add_run()
                path = item["item"].mImgPath
                width = item["item"].mWidth
                kh.add_picture(path, width=Inches(2))

    # Escribe texto en el body
    def writeText(self, text: textDocx):
        self.generateText(text, self.myTextDoc)

    # Escribe imagen en el body
    def writeImage(self, image: imageDocx):
        self.generateImage(image, self.myTextDoc)

    # Escribe tabla en el body
    def writeTable(self, table: tableDocx):
        columns = int(table.mColumns)
        rows = int(table.mRows)
        table.sortItems()
        docTable = self.myTextDoc.add_table(rows=rows, cols=columns, style='Table Grid')
        for item in table.mItems:
            celda = docTable.cell(int(item.mRow), int(item.mColumn))
            if not item.mVoid:
                #self.generateText(item,celda.paragraphs[0])
                paragraph = celda.paragraphs[0]
                text = item.mItem

                # Establecer el contenido del párrafo
                content = text.mContent
                paragraph.add_run(content)

                # Aplicar formatos
                run = paragraph.runs[0]
                if text.mBold == "1":
                    run.bold = True
                if text.mUnderline == "1":
                    run.underline = True
                if text.mItalic == "1":
                    run.italic = True
                if text.mStrikethrough == "1":
                    run.strike = True

                # Establecer alineación
                alignment = text.mAlign.lower()
                paragraph.alignment = Document().styles['Normal'].paragraph_format.alignment
                if alignment == 'right':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif alignment == 'center':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == 'justify':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                # Establecer color de fuente
                rgb = separarRgb(text.mFillColor)
                run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

                # Establecer color de fondo
                # rgb = separarRgb(text.mBackgroundColor)
                # paragraph.runs[0].font.highlight_color = RGBColor(rgb[0], rgb[1], rgb[2])

                # Establecer fuente y tamaño
                run.font.name = text.mFont
                run.font.size = Pt(int(text.mLetterSize))

            # Alinear verticalmente el contenido de la celda
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Escribe lista en el body
    def writeList(self, textList: listDocx):
        lista = textList.mListItems
        for text in lista:
            paragraph = self.myTextDoc.add_paragraph()
            paragraph.style = 'List Bullet'

            # Establecer el contenido del párrafo
            content = text.mContent
            paragraph.add_run(content)

            # Aplicar formatos
            run = paragraph.runs[0]
            if text.mBold == "1":
                run.bold = True
            if text.mUnderline == "1":
                run.underline = True
            if text.mItalic == "1":
                run.italic = True
            if text.mStrikethrough == "1":
                run.strike = True

            # Establecer alineación
            alignment = text.mAlign.lower()
            paragraph.alignment = Document().styles['Normal'].paragraph_format.alignment
            if alignment == 'right':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == 'center':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 'justify':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Establecer color de fuente
            rgb = separarRgb(text.mFillColor)
            run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

            # Establecer color de fondo
            # rgb = separarRgb(text.mBackgroundColor)
            # paragraph.runs[0].font.highlight_color = RGBColor(rgb[0], rgb[1], rgb[2])

            # Establecer fuente y tamaño
            run.font.name = text.mFont
            run.font.size = Pt(int(text.mLetterSize))

    # Guarda y cierra el archivo
    def closeFile(self):
        self.myTextDoc.save(self.path)
